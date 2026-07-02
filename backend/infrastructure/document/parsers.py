"""Multi-format local document parsers — zero network calls."""
from __future__ import annotations

import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# ── MIME detection ─────────────────────────────────────────────────────────────

_MIME_BY_EXT: dict[str, str] = {
    ".txt": "text/plain",
    ".md": "text/markdown",
    ".rst": "text/x-rst",
    ".csv": "text/csv",
    ".json": "application/json",
    ".xml": "application/xml",
    ".html": "text/html",
    ".htm": "text/html",
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    ".py": "text/x-python",
    ".js": "text/javascript",
    ".ts": "text/typescript",
    ".java": "text/x-java",
    ".c": "text/x-c",
    ".cpp": "text/x-c++",
    ".go": "text/x-go",
    ".rs": "text/x-rust",
}


def detect_mime(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    return _MIME_BY_EXT.get(ext, "application/octet-stream")


# ── Base parser ────────────────────────────────────────────────────────────────

class ParseResult:
    def __init__(self, text: str, metadata: dict | None = None):
        self.text = text
        self.metadata: dict = metadata or {}


class BaseParser:
    supported_mimes: list[str] = []

    def can_parse(self, mime: str) -> bool:
        return mime in self.supported_mimes

    def parse(self, data: bytes, filename: str) -> ParseResult:
        raise NotImplementedError


# ── Plain-text parser ──────────────────────────────────────────────────────────

class PlainTextParser(BaseParser):
    supported_mimes = [
        "text/plain", "text/markdown", "text/x-rst", "text/csv",
        "text/x-python", "text/javascript", "text/typescript",
        "text/x-java", "text/x-c", "text/x-c++", "text/x-go", "text/x-rust",
    ]

    def parse(self, data: bytes, filename: str) -> ParseResult:
        for enc in ("utf-8", "latin-1", "cp1252"):
            try:
                text = data.decode(enc)
                return ParseResult(text, {"encoding": enc})
            except UnicodeDecodeError:
                continue
        return ParseResult(data.decode("utf-8", errors="replace"), {"encoding": "utf-8-replace"})


# ── JSON parser ────────────────────────────────────────────────────────────────

class JSONParser(BaseParser):
    supported_mimes = ["application/json"]

    def parse(self, data: bytes, filename: str) -> ParseResult:
        import json
        try:
            obj = json.loads(data)
            text = json.dumps(obj, ensure_ascii=False, indent=2)
            return ParseResult(text, {"json_keys": list(obj.keys()) if isinstance(obj, dict) else []})
        except Exception as exc:
            logger.warning("JSON parse failed for %s: %s", filename, exc)
            return ParseResult(data.decode("utf-8", errors="replace"))


# ── HTML parser ────────────────────────────────────────────────────────────────

class HTMLParser(BaseParser):
    supported_mimes = ["text/html"]

    def parse(self, data: bytes, filename: str) -> ParseResult:
        try:
            from html.parser import HTMLParser as StdHTMLParser

            class _Extractor(StdHTMLParser):
                def __init__(self):
                    super().__init__()
                    self._parts: list[str] = []
                    self._skip = False

                def handle_starttag(self, tag, attrs):
                    if tag in ("script", "style"):
                        self._skip = True

                def handle_endtag(self, tag):
                    if tag in ("script", "style"):
                        self._skip = False

                def handle_data(self, data):
                    if not self._skip:
                        stripped = data.strip()
                        if stripped:
                            self._parts.append(stripped)

            p = _Extractor()
            p.feed(data.decode("utf-8", errors="replace"))
            return ParseResult(" ".join(p._parts))
        except Exception as exc:
            logger.warning("HTML parse failed for %s: %s", filename, exc)
            return ParseResult(data.decode("utf-8", errors="replace"))


# ── PDF parser (pdfminer.six — pure Python, no binary deps) ───────────────────

class PDFParser(BaseParser):
    supported_mimes = ["application/pdf"]

    def parse(self, data: bytes, filename: str) -> ParseResult:
        try:
            from pdfminer.high_level import extract_text  # type: ignore
            text = extract_text(io.BytesIO(data))
            return ParseResult(text or "", {"source": "pdfminer"})
        except ImportError:
            logger.warning("pdfminer.six not installed; PDF will be skipped.")
            return ParseResult("", {"error": "pdfminer not available"})
        except Exception as exc:
            logger.warning("PDF parse failed for %s: %s", filename, exc)
            return ParseResult("", {"error": str(exc)})


# ── DOCX parser (python-docx) ──────────────────────────────────────────────────

class DOCXParser(BaseParser):
    supported_mimes = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ]

    def parse(self, data: bytes, filename: str) -> ParseResult:
        try:
            import docx  # type: ignore
            doc = docx.Document(io.BytesIO(data))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return ParseResult(text, {"paragraphs": len(doc.paragraphs)})
        except ImportError:
            logger.warning("python-docx not installed; DOCX will be skipped.")
            return ParseResult("", {"error": "python-docx not available"})
        except Exception as exc:
            logger.warning("DOCX parse failed for %s: %s", filename, exc)
            return ParseResult("", {"error": str(exc)})


# ── XLSX parser (openpyxl) ─────────────────────────────────────────────────────

class XLSXParser(BaseParser):
    supported_mimes = [
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    ]

    def parse(self, data: bytes, filename: str) -> ParseResult:
        try:
            import openpyxl  # type: ignore
            wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
            parts: list[str] = []
            for sheet in wb.worksheets:
                parts.append(f"=== Sheet: {sheet.title} ===")
                for row in sheet.iter_rows(values_only=True):
                    cells = [str(c) for c in row if c is not None]
                    if cells:
                        parts.append("\t".join(cells))
            return ParseResult("\n".join(parts), {"sheets": wb.sheetnames})
        except ImportError:
            logger.warning("openpyxl not installed; XLSX will be skipped.")
            return ParseResult("", {"error": "openpyxl not available"})
        except Exception as exc:
            logger.warning("XLSX parse failed for %s: %s", filename, exc)
            return ParseResult("", {"error": str(exc)})


# ── Registry ───────────────────────────────────────────────────────────────────

_PARSERS: list[BaseParser] = [
    PDFParser(),
    DOCXParser(),
    XLSXParser(),
    HTMLParser(),
    JSONParser(),
    PlainTextParser(),
]


def get_parser(mime: str) -> BaseParser | None:
    for p in _PARSERS:
        if p.can_parse(mime):
            return p
    return None
