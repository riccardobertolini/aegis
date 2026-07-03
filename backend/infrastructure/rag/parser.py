"""DocumentParser — extracts plain text from PDF, DOCX, TXT, MD, CSV.

All parsing is 100% local. No network calls. No cloud OCR.
Fallback: raw UTF-8 decode with error replacement.
"""
from __future__ import annotations

import hashlib
import io
import logging
from pathlib import Path

from backend.domain.ports.document import ParsedDocument

logger = logging.getLogger(__name__)

# MIME type constants
_MIME_PDF = "application/pdf"
_MIME_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_MIME_DOC = "application/msword"
_MIME_TXT = "text/plain"
_MIME_MD = "text/markdown"
_MIME_CSV = "text/csv"
_MIME_HTML = "text/html"


class DocumentParser:
    """Parses documents into plain-text chunks.

    Supported formats:
        - PDF  (pdfminer.six)
        - DOCX (python-docx)
        - TXT / MD / CSV / HTML — direct UTF-8 decode
        - Fallback: bytes decoded as UTF-8 with error replacement
    """

    def __init__(self, max_file_bytes: int = 100 * 1024 * 1024) -> None:
        self._max_file_bytes = max_file_bytes

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def parse_file(self, path: Path) -> ParsedDocument:
        """Parse a file from disk."""
        if path.stat().st_size > self._max_file_bytes:
            raise ValueError(
                f"File {path.name} exceeds max size "
                f"({path.stat().st_size} > {self._max_file_bytes} bytes)"
            )
        data = path.read_bytes()
        mime = self._detect_mime(path.name, data)
        return self._parse(data, path.name, mime)

    def parse_bytes(self, data: bytes, filename: str) -> ParsedDocument:
        """Parse raw bytes (e.g. from HTTP upload)."""
        if len(data) > self._max_file_bytes:
            raise ValueError(
                f"Upload {filename} exceeds max size "
                f"({len(data)} > {self._max_file_bytes} bytes)"
            )
        mime = self._detect_mime(filename, data)
        return self._parse(data, filename, mime)

    # ------------------------------------------------------------------
    # Internal dispatch
    # ------------------------------------------------------------------

    def _parse(self, data: bytes, filename: str, mime: str) -> ParsedDocument:
        doc_id = self._make_id(data, filename)
        try:
            if mime == _MIME_PDF:
                text = self._parse_pdf(data)
            elif mime in (_MIME_DOCX, _MIME_DOC):
                text = self._parse_docx(data)
            elif mime in (_MIME_TXT, _MIME_MD, _MIME_CSV, _MIME_HTML):
                text = self._parse_text(data)
            else:
                text = self._parse_fallback(data, filename)
        except Exception as exc:
            logger.warning("Parser failed for %s (%s): %s — using fallback", filename, mime, exc)
            text = self._parse_fallback(data, filename)

        # Split into raw text chunks (one per page / paragraph block)
        chunks = self._split_to_chunks(text)
        logger.info("Parsed '%s' → %d chunk(s) (mime=%s)", filename, len(chunks), mime)

        return ParsedDocument(
            id=doc_id,
            filename=filename,
            mime_type=mime,
            chunks=chunks,
            metadata={
                "size_bytes": len(data),
                "checksum_sha256": hashlib.sha256(data).hexdigest(),
                "mime_type": mime,
            },
        )

    # ------------------------------------------------------------------
    # Format-specific parsers
    # ------------------------------------------------------------------

    @staticmethod
    def _parse_pdf(data: bytes) -> str:
        try:
            from pdfminer.high_level import extract_text  # type: ignore
            return extract_text(io.BytesIO(data)) or ""
        except ImportError as exc:
            raise ImportError(
                "pdfminer.six is required for PDF parsing. "
                "Add it to requirements/base.txt."
            ) from exc

    @staticmethod
    def _parse_docx(data: bytes) -> str:
        try:
            import docx  # type: ignore
            doc = docx.Document(io.BytesIO(data))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError as exc:
            raise ImportError(
                "python-docx is required for DOCX parsing. "
                "Add it to requirements/base.txt."
            ) from exc

    @staticmethod
    def _parse_text(data: bytes) -> str:
        return data.decode("utf-8", errors="replace")

    @staticmethod
    def _parse_fallback(data: bytes, filename: str) -> str:
        logger.warning("Using raw-bytes fallback for '%s'", filename)
        return data.decode("utf-8", errors="replace")

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _detect_mime(filename: str, data: bytes) -> str:
        # Magic bytes detection first
        if data[:4] == b"%PDF":
            return _MIME_PDF
        if data[:4] == b"PK\x03\x04":  # ZIP-based formats (DOCX)
            return _MIME_DOCX
        if data[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":  # OLE2 (DOC)
            return _MIME_DOC

        # Extension fallback
        ext = Path(filename).suffix.lower()
        ext_map = {
            ".pdf": _MIME_PDF,
            ".docx": _MIME_DOCX,
            ".doc": _MIME_DOC,
            ".txt": _MIME_TXT,
            ".md": _MIME_MD,
            ".csv": _MIME_CSV,
            ".html": _MIME_HTML,
            ".htm": _MIME_HTML,
        }
        return ext_map.get(ext, _MIME_TXT)

    @staticmethod
    def _make_id(data: bytes, filename: str) -> str:
        digest = hashlib.sha256(data + filename.encode()).hexdigest()[:16]
        return f"doc_{digest}"

    @staticmethod
    def _split_to_chunks(text: str) -> list[str]:
        """Split text into paragraph-level chunks (double-newline boundaries)."""
        raw = [c.strip() for c in text.split("\n\n") if c.strip()]
        # Merge very short fragments into the previous chunk
        merged: list[str] = []
        for block in raw:
            if merged and len(block) < 80:
                merged[-1] += " " + block
            else:
                merged.append(block)
        return merged or [text.strip()] if text.strip() else []
