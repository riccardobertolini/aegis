"""DOCX parser using python-docx (offline)."""
from pathlib import Path

from ..models import DocumentFormat, ParsedDocument
from .base import BaseParser


class DocxParser(BaseParser):
    supported_formats = [DocumentFormat.DOCX]

    def parse(self, path: Path) -> ParsedDocument:
        try:
            import docx

            doc = docx.Document(str(path))
            raw_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            props = doc.core_properties
            author = props.author or ""
            title = props.title or path.stem
            created_at = props.created
            modified_at = props.modified
        except ImportError:
            raw_text = f"[python-docx not installed — cannot parse {path.name}]"
            author, title, created_at, modified_at = "", path.stem, None, None
        except Exception as exc:  # noqa: BLE001
            raw_text = f"[DOCX parse error: {exc}]"
            author, title, created_at, modified_at = "", path.stem, None, None

        return ParsedDocument(
            source_path=str(path),
            format=DocumentFormat.DOCX,
            title=title,
            author=author,
            created_at=created_at,
            modified_at=modified_at,
            raw_text=raw_text,
        )
